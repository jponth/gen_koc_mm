from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class MinutesJSON:
    date_of_meeting: str
    sections: dict[str, str]  # section_key -> section_text (markdown)


def load_minutes_json(path: Path) -> MinutesJSON:
    payload = json.loads(path.read_text(encoding="utf-8"))

    date = (payload.get("date_of_meeting") or "").strip()

    sections_map: dict[str, str] = {}
    for s in payload.get("sections") or []:
        key = (s.get("section_key") or "").strip()
        txt = (s.get("section_text") or "").rstrip()
        if key:
            sections_map[key] = txt

    return MinutesJSON(date_of_meeting=date, sections=sections_map)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    """Insert a new paragraph after the given paragraph."""

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def _clear_paragraph(paragraph: Paragraph) -> None:
    for r in paragraph.runs[::-1]:
        r._element.getparent().remove(r._element)


def _replace_placeholder_in_paragraph(
    paragraph: Paragraph,
    placeholder: str,
    replacement: str,
    *,
    bullet_style: str | None = None,
) -> bool:
    """Replace placeholder in a paragraph.

    Supports a special case where the paragraph is *just* the placeholder, in
    which case we can expand markdown-ish bullets ("- ") into true Word bullet
    paragraphs.

    Returns True if a replacement was applied.
    """

    if placeholder not in paragraph.text:
        return False

    # Special case: paragraph contains only placeholder (common for templates).
    if paragraph.text.strip() == placeholder:
        lines = replacement.splitlines()
        # If it's bullet-like markdown, convert to real Word bullet paragraphs.
        if any(ln.strip().startswith("- ") for ln in lines):
            # We'll reuse the existing paragraph as the first output paragraph.
            cur = paragraph
            first = True

            for ln in lines:
                stripped = ln.rstrip()
                if not stripped:
                    # Blank line → empty normal paragraph
                    cur = paragraph if first else _insert_paragraph_after(cur, "")
                    cur.style = "Normal"
                    _clear_paragraph(cur)
                    first = False
                    continue

                is_bullet = stripped.lstrip().startswith("- ")
                text = stripped.lstrip()[2:] if is_bullet else stripped
                if is_bullet and bullet_style:
                    style = bullet_style
                elif is_bullet:
                    # Fallback: insert a literal bullet if the template has no bullet styles.
                    style = "Normal"
                    text = f"• {text}" if text else "•"
                else:
                    style = "Normal"

                if first:
                    _clear_paragraph(cur)
                    cur.style = style
                    cur.add_run(text)
                    first = False
                else:
                    cur = _insert_paragraph_after(cur, text=text, style=style)

            return True

        # Non-bullet multi-line content: keep as literal line breaks.
        _clear_paragraph(paragraph)
        lines2 = replacement.splitlines() or [""]
        run = paragraph.add_run(lines2[0])
        for ln in lines2[1:]:
            run.add_break()
            paragraph.add_run(ln)
        return True

    # Fallback: paragraph has other text; do a literal replacement.
    new_text = paragraph.text.replace(placeholder, replacement)
    _clear_paragraph(paragraph)
    lines = new_text.splitlines() or [""]
    run = paragraph.add_run(lines[0])
    for ln in lines[1:]:
        run.add_break()
        paragraph.add_run(ln)
    return True


def _iter_all_paragraphs(doc: Document):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p

    # Table cells paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _pick_bullet_style_name(doc: Document) -> str | None:
    # Different templates can have different style sets.
    # Try a few common bullet paragraph styles.
    candidates = [
        "List Bullet",
        "ListBullet",
        "Bullet List",
        "Bulleted List",
        "List Paragraph",
        "ListParagraph",
    ]
    available = {s.name for s in doc.styles}
    for c in candidates:
        if c in available:
            return c
    return None


def _format_date_mmm_dd_yyyy(date_str: str) -> str:
    """Format YYYY-MM-DD -> 'Mon DD, YYYY'.

    If parsing fails, returns the original string.
    """
    raw = (date_str or "").strip()
    if not raw:
        return ""

    # Accept strict ISO date first.
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try:
            dt = datetime.strptime(raw, fmt)
            return dt.strftime("%b %d, %Y")
        except ValueError:
            pass

    # If we got a timestamp-ish string like 'YYYY-MM-DD ...', take the date part.
    if len(raw) >= 10 and raw[4:5] == "-" and raw[7:8] == "-":
        try:
            dt = datetime.strptime(raw[:10], "%Y-%m-%d")
            return dt.strftime("%b %d, %Y")
        except ValueError:
            pass

    return raw


def merge_minutes_into_docx(*, minutes_json_path: Path, template_docx_path: Path, output_docx_path: Path) -> dict[str, Any]:
    minutes = load_minutes_json(minutes_json_path)

    doc = Document(str(template_docx_path))
    bullet_style = _pick_bullet_style_name(doc)

    replaced: list[str] = []

    # Special top-level placeholder (optional)
    date_placeholder = "<<date_of_meeting>>"
    formatted_date = _format_date_mmm_dd_yyyy(minutes.date_of_meeting)
    for p in _iter_all_paragraphs(doc):
        if _replace_placeholder_in_paragraph(p, date_placeholder, formatted_date, bullet_style=bullet_style):
            replaced.append(date_placeholder)

    for key, txt in minutes.sections.items():
        ph = f"<<{key}>>"
        for p in _iter_all_paragraphs(doc):
            if _replace_placeholder_in_paragraph(p, ph, txt, bullet_style=bullet_style):
                replaced.append(ph)

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))

    return {
        "output": str(output_docx_path),
        "replaced_placeholders": sorted(set(replaced)),
    }
