from __future__ import annotations

from pathlib import Path

from docx import Document


def make_template(*, output_path: Path) -> None:
    doc = Document()

    doc.add_heading("Knights of Columbus – Council Meeting Minutes", level=1)

    p = doc.add_paragraph()
    p.add_run("Date of meeting: ")
    p.add_run("<<date_of_meeting>>")

    doc.add_paragraph("")

    sections = [
        ("Chaplain’s Report", "chaplains_report"),
        ("Grand Knights Report", "grand_knights_report"),
        ("Treasurer’s Report", "treasurers_report"),
        ("Financial Secretary’s Report", "financial_secretary_report"),
        ("Church Director's Report", "church_director_report"),
        ("Insurance Agent Report", "insurance_agent_report"),
        ("Social Action Report", "social_action_report"),
        ("District Deputy Report", "district_deputy_report"),
        ("4th Degree Report", "4th_degree_report"),
        ("Old Business", "old_business"),
        ("New Business", "new_business"),
        ("Birthdays", "birthdays"),
        ("Good of the Order", "good_of_the_order"),
        ("Closing Prayers", "closing_prayers"),
    ]

    for heading, key in sections:
        doc.add_heading(heading, level=2)
        # Placeholder must be the only content in the paragraph so merge-docx
        # can expand markdown bullets into true Word bullet paragraphs.
        doc.add_paragraph(f"<<{key}>>")
        doc.add_paragraph("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "templates" / "KoC-Meeting-Minutes-Template-compatible.docx"
    make_template(output_path=out)
    print(f"Wrote template: {out}")
